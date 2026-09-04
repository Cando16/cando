import pytest
from fastapi.testclient import TestClient
from backend.main import app
import docx
import io

client = TestClient(app)

def test_process_docx_endpoint():
    # Create an in-memory docx
    doc = docx.Document()
    doc.add_paragraph("Furthermore, it is important to note that the sky is blue.")
    
    in_stream = io.BytesIO()
    doc.save(in_stream)
    in_stream.seek(0)
    
    response = client.post(
        "/api/file/process",
        files={"file": ("test.docx", in_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    out_doc = docx.Document(io.BytesIO(response.content))
    # We expect 'Furthermore, it is important to note that' to become 'Additionally, ' via No-AI-Slop
    assert "Additionally, the sky is blue." in out_doc.paragraphs[0].text
