import io
import docx
from pypdf import PdfReader
from backend.core import pipeline, humanizer, metadata

async def process_docx(file_bytes: bytes, options: humanizer.CandoOptions) -> bytes:
    doc = docx.Document(io.BytesIO(file_bytes))
    
    for p in doc.paragraphs:
        if p.text.strip():
            result = await pipeline.run_cando(p.text, options)
            p.text = result.result
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        result = await pipeline.run_cando(p.text, options)
                        p.text = result.result

    metadata.scrub_docx_metadata(doc)
    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()

async def process_pdf(file_bytes: bytes, options: humanizer.CandoOptions) -> bytes:
    reader = PdfReader(io.BytesIO(file_bytes))
    
    new_doc = docx.Document()
    new_doc.add_heading("CANDO Extracted PDF", 0)
    
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            result = await pipeline.run_cando(text, options)
            new_doc.add_paragraph(result.result)
            
    metadata.scrub_docx_metadata(new_doc)
    out_stream = io.BytesIO()
    new_doc.save(out_stream)
    return out_stream.getvalue()
